# tools/file_handler.py

import json
import logging
import platform
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd

log = logging.getLogger(__name__)

MAX_TEXT_OUTPUT = 2000
MAX_JSON_OUTPUT = 1000


def detect_file_type(filepath: str) -> str:
    suffix = Path(filepath).suffix.lower()
    type_map = {
        ".txt": "text", ".md": "text", ".py": "text", ".html": "text",
        ".csv": "csv",
        ".pdf": "pdf",
        ".json": "json",
        ".xlsx": "excel", ".xls": "excel",
        ".docx": "docx", ".doc": "docx"
    }
    return type_map.get(suffix, "unknown")


def get_file_metadata(filepath: str) -> str:
    try:
        path = Path(filepath)
        if not path.exists():
            return f"File not found: {filepath}"
        stat = path.stat()
        size_kb = round(stat.st_size / 1024, 2)
        created = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S")
        modified = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        return (
            f"Name     : {path.name}\n"
            f"Type     : {detect_file_type(filepath)}\n"
            f"Size     : {size_kb} KB\n"
            f"Created  : {created}\n"
            f"Modified : {modified}\n"
            f"Path     : {path.resolve()}"
        )
    except Exception as e:
        log.error(f"Metadata extraction failed: {e}")
        return f"Metadata error: {e}"


def read_txt(filepath: str) -> str:
    try:
        content = Path(filepath).read_text(encoding="utf-8")
        log.info(f"TXT read — {len(content)} characters")
        return content[:MAX_TEXT_OUTPUT]
    except Exception as e:
        log.error(f"TXT read failed: {e}")
        return f"TXT read error: {e}"


def read_csv(filepath: str) -> str:
    try:
        df = pd.read_csv(filepath)
        stats = df.describe(include="all").to_string()
        result = (
            f"Shape    : {df.shape}\n"
            f"Columns  : {list(df.columns)}\n\n"
            f"Preview:\n{df.head(5).to_string()}\n\n"
            f"Stats:\n{stats[:MAX_TEXT_OUTPUT]}"
        )
        log.info(f"CSV read — shape: {df.shape}")
        return result
    except Exception as e:
        log.error(f"CSV read failed: {e}")
        return f"CSV read error: {e}"


def read_pdf(filepath: str) -> str:
    try:
        import PyPDF2
        text = ""
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text += f"\n--- Page {i+1} ---\n"
                text += page.extract_text() or ""
        log.info(f"PDF read — {len(reader.pages)} pages")
        return text[:MAX_TEXT_OUTPUT]
    except ImportError:
        return "PyPDF2 not installed. Run: pip install PyPDF2"
    except Exception as e:
        log.error(f"PDF read failed: {e}")
        return f"PDF read error: {e}"


def read_json(filepath: str) -> str:
    try:
        data = json.loads(Path(filepath).read_text(encoding="utf-8"))
        output = json.dumps(data, indent=2)
        log.info(f"JSON read — {len(output)} characters")
        return output[:MAX_JSON_OUTPUT]
    except Exception as e:
        log.error(f"JSON read failed: {e}")
        return f"JSON read error: {e}"


def read_excel(filepath: str) -> str:
    try:
        df = pd.read_excel(filepath)
        result = (
            f"Shape   : {df.shape}\n"
            f"Columns : {list(df.columns)}\n\n"
            f"Preview:\n{df.head(5).to_string()}"
        )
        log.info(f"Excel read — shape: {df.shape}")
        return result
    except Exception as e:
        log.error(f"Excel read failed: {e}")
        return f"Excel read error: {e}"

def read_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)

        content = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    content.append(f"\n**{para.text.strip()}**")
                else:
                    content.append(para.text.strip())

        # Extract tables
        for i, table in enumerate(doc.tables):
            content.append(f"\n**Table {i+1}:**")
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip():
                    content.append(row_text)

        result = "\n".join(content)
        log.info(f"DOCX read — {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        return result[:MAX_TEXT_OUTPUT] if result else "Document appears to be empty."

    except Exception as e:
        log.error(f"DOCX read failed: {e}")
        return f"DOCX read error: {e}"


def scan_directory(dirpath: str) -> str:
    try:
        path = Path(dirpath)
        if not path.exists():
            return f"Directory not found: {dirpath}"
        files = list(path.rglob("*"))
        lines = [f"Directory: {path.resolve()} — {len(files)} items\n"]
        for f in sorted(files):
            if f.is_file():
                size_kb = round(f.stat().st_size / 1024, 2)
                lines.append(f"  {f.name} ({detect_file_type(str(f))}) — {size_kb} KB")
        log.info(f"Directory scanned — {len(files)} items")
        return "\n".join(lines)
    except Exception as e:
        log.error(f"Directory scan failed: {e}")
        return f"Directory scan error: {e}"


def read_file(filepath: str) -> str:
    path = Path(filepath)
    if not path.exists():
        return f"File not found: {filepath}"
    file_type = detect_file_type(filepath)
    log.info(f"read_file called — type: {file_type} | path: {filepath}")
    routes = {
        "text": read_txt,
        "csv": read_csv,
        "pdf": read_pdf,
        "json": read_json,
        "excel": read_excel
        "docx": read_docx
    }
    handler = routes.get(file_type)
    if not handler:
        return f"Unsupported file type: {file_type}"
    return handler(filepath)