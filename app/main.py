# app/main.py

import sys
import logging
import json
from pathlib import Path
from typing import Optional
from datetime import datetime

sys.path.append(str(Path(__file__).resolve().parent.parent))

# Added UploadFile and File for direct data loading
from fastapi import FastAPI, HTTPException, UploadFile, File as FastAPIFile
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from agent.core import AgentConfig, init_client, AGENT_CAPABILITIES, EXPORT_PATH, BASE_PATH
from memory.short_term import ShortTermBuffer
from memory.long_term import LongTermMemory
from tools.web_search import SearchHistoryManager
from tools.calendar_tool import CalendarManager
from agent.orchestrator import orchestrate
from agent.response_builder import (
    build_startup_banner, build_help_text,
    build_session_summary, build_stats_text, build_response
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# --- tool registry ---
import math
import requests
import platform
import psutil
import shutil  # Added for direct background file stream copying
import pandas as pd
import wikipediaapi
from ddgs import DDGS

MAX_TOOL_OUTPUT = 500

def calculator(expression: str) -> str:
    import re
    try:
        clean = expression.lower().strip()
        
        # Extract variable assignments like "x is 5" or "x = 5"
        var_pattern = re.findall(r'([a-zA-Z])\s*(?:is|=)\s*([\d.]+)', clean)
        variables = {var: float(val) for var, val in var_pattern}

        # Substitute variables into expression
        for var, val in variables.items():
            clean = re.sub(rf'\b{var}\b', str(val), clean)

        # Handle percentage
        percent_match = re.search(r'(\d+\.?\d*)\s*%\s*of\s*(\d+\.?\d*)', clean)
        if percent_match:
            percent = float(percent_match.group(1))
            total = float(percent_match.group(2))
            result = (percent / 100) * total
            return f"Result: {round(result, 6)}"

        # Handle natural language math
        clean = re.sub(r'\bplus\b', '+', clean)
        clean = re.sub(r'\bminus\b', '-', clean)
        clean = re.sub(r'\btimes\b', '*', clean)
        clean = re.sub(r'\bdivided by\b', '/', clean)
        clean = re.sub(r'\bto the power of\b', '**', clean)
        clean = re.sub(r'\bsquared\b', '**2', clean)
        clean = re.sub(r'\bcubed\b', '**3', clean)
        clean = re.sub(r'\bsquare root of\b', 'math.sqrt', clean)

        # Handle degree conversion
        has_degrees = 'degree' in clean or '°' in clean
        clean = re.sub(r'\bdegrees?\b', '', clean)
        clean = re.sub(r'°', '', clean)

        # Handle trig with ^ notation
        clean = re.sub(r'sin\^2\(([^)]+)\)', r'math.sin(\1)**2', clean)
        clean = re.sub(r'cos\^2\(([^)]+)\)', r'math.cos(\1)**2', clean)
        clean = re.sub(r'tan\^2\(([^)]+)\)', r'math.tan(\1)**2', clean)

        # Handle trig functions
        if has_degrees:
            clean = re.sub(r'sin\(([^)]+)\)', r'math.sin(math.radians(\1))', clean)
            clean = re.sub(r'cos\(([^)]+)\)', r'math.cos(math.radians(\1))', clean)
            clean = re.sub(r'tan\(([^)]+)\)', r'math.tan(math.radians(\1))', clean)
        else:
            clean = re.sub(r'sin\(([^)]+)\)', r'math.sin(\1)', clean)
            clean = re.sub(r'cos\(([^)]+)\)', r'math.cos(\1)', clean)
            clean = re.sub(r'tan\(([^)]+)\)', r'math.tan(\1)', clean)

        # Handle log functions
        clean = re.sub(r'log10\(([^)]+)\)', r'math.log10(\1)', clean)
        clean = re.sub(r'log2\(([^)]+)\)', r'math.log2(\1)', clean)
        clean = re.sub(r'log\(([^)]+)\)', r'math.log(\1)', clean)

        # Handle other math functions
        clean = re.sub(r'sqrt\(([^)]+)\)', r'math.sqrt(\1)', clean)
        clean = re.sub(r'abs\(([^)]+)\)', r'abs(\1)', clean)
        clean = re.sub(r'ceil\(([^)]+)\)', r'math.ceil(\1)', clean)
        clean = re.sub(r'floor\(([^)]+)\)', r'math.floor(\1)', clean)
        clean = re.sub(r'factorial\(([^)]+)\)', r'math.factorial(int(\1))', clean)
        clean = re.sub(r'round\(([^)]+)\)', r'round(\1)', clean)

        # Handle constants
        clean = re.sub(r'\bpi\b', 'math.pi', clean)
        clean = re.sub(r'\beuler\b|\b\be\b', 'math.e', clean)

        # Handle ^ as power
        clean = re.sub(r'\^', '**', clean)

        # Extract only the math expression
        match = re.search(r'[\d\s\+\-\*\/\(\)\.\%math\.piesqrtlogsincotan]+', clean)
        expr = match.group(0).strip() if match else clean

        safe_dict = {
            "__builtins__": {},
            "math": math,
            "abs": abs,
            "round": round,
            "int": int,
            "float": float
        }

        result = eval(expr, safe_dict)
        return f"Result: {round(float(result), 6)}"

    except Exception as e:
        return f"Calculator error: {e}. Please use format like '25 * 48' or 'sin(30 degrees)'"

def wikipedia_tool(query: str) -> str:
    import re
    try:
        clean_query = re.sub(
            r'(search wikipedia for|wikipedia|tell me about|who is|what is|explain)',
            '', query.lower()
        ).strip()

        wiki = wikipediaapi.Wikipedia(
            language="en",
            user_agent="AutonomousAgent/1.0"
        )

        page = wiki.page(clean_query)

        if not page.exists():
            # Try title case
            page = wiki.page(clean_query.title())

        if not page.exists():
            return f"No Wikipedia page found for: {clean_query}"

        # Get summary
        summary = page.summary[:800].strip()

        # Get top sections
        sections = []
        for section in list(page.sections)[:3]:
            if section.text:
                sections.append(f"**{section.title}**\n{section.text[:200].strip()}")

        # Build response
        result = f"**{page.title}**\n\n"
        result += f"{summary}\n\n"

        if sections:
            result += "**Key Sections:**\n"
            result += "\n\n".join(sections)

        result += f"\n\nSource: {page.fullurl}"

        return result

    except Exception as e:
        return f"Wikipedia error: {e}"

def web_search_tool(query: str) -> str:
    import re
    try:
        news_keywords = [
            'news', 'latest', 'today', 'recent', 'breaking',
            'update', 'current', 'happening', 'just in', 'announced'
        ]
        query_lower = query.lower()
        is_news = any(keyword in query_lower for keyword in news_keywords)

        with DDGS() as ddgs:
            if is_news:
                results = list(ddgs.news(query, max_results=5))
                if results:
                    lines = [f"**Latest News: {query}**\n"]
                    for i, r in enumerate(results, 1):
                        title = r.get('title', '')
                        source = r.get('source', '')
                        date = r.get('date', '')
                        body = r.get('body', '')[:150]
                        url = r.get('url', '')
                        lines.append(
                            f"{i}. **{title}**\n"
                            f"   Source: {source} | {date}\n"
                            f"   {body}\n"
                            f"   {url}\n"
                        )
                    return "\n".join(lines)

            # Regular web search
            results = list(ddgs.text(query, max_results=3))
            if results:
                return "\n".join([
                    f"{r['title']}: {r['body'][:150]}"
                    for r in results
                ])
            return "No results found."

    except Exception as e:
        return f"Web search error: {e}"
        
def file_reader(filepath: str) -> str:
    import re
    from pathlib import Path

    match = re.search(r'([A-Za-z]:[^\|]+\.\w+|/[^\|]+\.\w+)', filepath)
    actual_path = match.group(1).strip() if match else filepath.strip()
    actual_path = actual_path.replace("\\", "/")
    path = Path(actual_path)

    if not path.exists():
        return f"Error: The file at path {actual_path} was not found on the server disk storage."

    if path.suffix.lower() == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text_content = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            extracted_text = "\n".join(text_content).strip()
            if not extracted_text:
                return "Error: The PDF file appears to be empty or contains only scanned image layers."
            return extracted_text[:8000]
        except Exception as pdf_err:
            return f"Failed to extract text from PDF: {str(pdf_err)}"

    if path.suffix.lower() in [".docx", ".doc"]:
        try:
            from docx import Document
            doc = Document(path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        content.append(row_text)
            result = "\n".join(content)
            return result[:8000] if result else "Document appears to be empty."
        except Exception as e:
            return f"DOCX read error: {e}"

    try:
        from tools.file_handler import read_file
        return read_file(actual_path)
    except Exception as e:
        return f"File reader error: {e}"

def datetime_tool(query: str) -> str:
    import pytz
    import re
    from datetime import datetime as dt

    TIMEZONE_MAP = {
        "india": "Asia/Kolkata", "ist": "Asia/Kolkata",
        "mumbai": "Asia/Kolkata", "delhi": "Asia/Kolkata",
        "hyderabad": "Asia/Kolkata", "bangalore": "Asia/Kolkata",
        "chennai": "Asia/Kolkata", "kolkata": "Asia/Kolkata",
        "pune": "Asia/Kolkata", "ahmedabad": "Asia/Kolkata",
        "london": "Europe/London", "uk": "Europe/London",
        "england": "Europe/London", "britain": "Europe/London",
        "new york": "America/New_York", "usa": "America/New_York",
        "america": "America/New_York", "washington": "America/New_York",
        "los angeles": "America/Los_Angeles", "california": "America/Los_Angeles",
        "chicago": "America/Chicago", "houston": "America/Chicago",
        "toronto": "America/Toronto", "canada": "America/Toronto",
        "vancouver": "America/Vancouver",
        "dubai": "Asia/Dubai", "uae": "Asia/Dubai",
        "abu dhabi": "Asia/Dubai",
        "singapore": "Asia/Singapore",
        "tokyo": "Asia/Tokyo", "japan": "Asia/Tokyo",
        "osaka": "Asia/Tokyo",
        "beijing": "Asia/Shanghai", "shanghai": "Asia/Shanghai",
        "china": "Asia/Shanghai", "hong kong": "Asia/Hong_Kong",
        "sydney": "Australia/Sydney", "australia": "Australia/Sydney",
        "melbourne": "Australia/Melbourne",
        "paris": "Europe/Paris", "france": "Europe/Paris",
        "germany": "Europe/Berlin", "berlin": "Europe/Berlin",
        "moscow": "Europe/Moscow", "russia": "Europe/Moscow",
        "pakistan": "Asia/Karachi", "karachi": "Asia/Karachi",
        "lahore": "Asia/Karachi",
        "bangladesh": "Asia/Dhaka", "dhaka": "Asia/Dhaka",
        "nepal": "Asia/Kathmandu", "kathmandu": "Asia/Kathmandu",
        "sri lanka": "Asia/Colombo", "colombo": "Asia/Colombo",
        "malaysia": "Asia/Kuala_Lumpur", "kuala lumpur": "Asia/Kuala_Lumpur",
        "indonesia": "Asia/Jakarta", "jakarta": "Asia/Jakarta",
        "thailand": "Asia/Bangkok", "bangkok": "Asia/Bangkok",
        "vietnam": "Asia/Ho_Chi_Minh", "ho chi minh": "Asia/Ho_Chi_Minh",
        "philippines": "Asia/Manila", "manila": "Asia/Manila",
        "south korea": "Asia/Seoul", "seoul": "Asia/Seoul",
        "egypt": "Africa/Cairo", "cairo": "Africa/Cairo",
        "nigeria": "Africa/Lagos", "lagos": "Africa/Lagos",
        "abuja": "Africa/Lagos",
        "south africa": "Africa/Johannesburg", "johannesburg": "Africa/Johannesburg",
        "kenya": "Africa/Nairobi", "nairobi": "Africa/Nairobi",
        "ghana": "Africa/Accra", "accra": "Africa/Accra",
        "ethiopia": "Africa/Addis_Ababa", "addis ababa": "Africa/Addis_Ababa",
        "brazil": "America/Sao_Paulo", "sao paulo": "America/Sao_Paulo",
        "mexico": "America/Mexico_City", "mexico city": "America/Mexico_City",
        "argentina": "America/Argentina/Buenos_Aires",
        "buenos aires": "America/Argentina/Buenos_Aires",
        "colombia": "America/Bogota", "bogota": "America/Bogota",
        "peru": "America/Lima", "lima": "America/Lima",
        "chile": "America/Santiago", "santiago": "America/Santiago",
        "spain": "Europe/Madrid", "madrid": "Europe/Madrid",
        "italy": "Europe/Rome", "rome": "Europe/Rome",
        "netherlands": "Europe/Amsterdam", "amsterdam": "Europe/Amsterdam",
        "switzerland": "Europe/Zurich", "zurich": "Europe/Zurich",
        "sweden": "Europe/Stockholm", "stockholm": "Europe/Stockholm",
        "norway": "Europe/Oslo", "oslo": "Europe/Oslo",
        "denmark": "Europe/Copenhagen", "copenhagen": "Europe/Copenhagen",
        "finland": "Europe/Helsinki", "helsinki": "Europe/Helsinki",
        "poland": "Europe/Warsaw", "warsaw": "Europe/Warsaw",
        "turkey": "Europe/Istanbul", "istanbul": "Europe/Istanbul",
        "israel": "Asia/Jerusalem", "jerusalem": "Asia/Jerusalem",
        "saudi arabia": "Asia/Riyadh", "riyadh": "Asia/Riyadh",
        "qatar": "Asia/Qatar", "doha": "Asia/Qatar",
        "kuwait": "Asia/Kuwait",
        "iraq": "Asia/Baghdad", "baghdad": "Asia/Baghdad",
        "iran": "Asia/Tehran", "tehran": "Asia/Tehran",
        "afghanistan": "Asia/Kabul", "kabul": "Asia/Kabul",
        "utc": "UTC", "gmt": "GMT",
        "new zealand": "Pacific/Auckland", "auckland": "Pacific/Auckland",
        "hawaii": "Pacific/Honolulu",
    }

    query_lower = query.lower()

    matched_tz_name = "Asia/Kolkata"
    matched_place = "India (IST)"

    for place, tz_name in TIMEZONE_MAP.items():
        if place in query_lower:
            matched_tz_name = tz_name
            matched_place = place.title()
            break

    try:
        tz = pytz.timezone(matched_tz_name)
        now = dt.now(tz)
    except Exception:
        tz = pytz.timezone("Asia/Kolkata")
        now = dt.now(tz)
        matched_place = "India (IST)"

    if "time" in query_lower or "what time" in query_lower:
        return f"Current time in {matched_place}: {now.strftime('%H:%M:%S %Z')}"
    if "day" in query_lower:
        return f"Today in {matched_place}: {now.strftime('%A, %Y-%m-%d')}"
    return f"Current datetime in {matched_place}: {now.strftime('%Y-%m-%d %H:%M:%S %Z')}"

def unit_converter(query: str) -> str:
    import re
    try:
        pattern = re.search(
            r'([\d.]+)\s*([a-zA-Z]+)\s+(?:to|in|into)\s+([a-zA-Z]+)',
            query.lower()
        )
        if not pattern:
            return f"Could not parse conversion request. Use format: '100 km to miles'"

        value = float(pattern.group(1))
        unit_from = pattern.group(2).strip()
        unit_to = pattern.group(3).strip()

        conversions = {
            ("kg", "lbs"): lambda x: x * 2.20462,
            ("lbs", "kg"): lambda x: x / 2.20462,
            ("km", "miles"): lambda x: x * 0.621371,
            ("miles", "km"): lambda x: x / 0.621371,
            ("kilometers", "miles"): lambda x: x * 0.621371,
            ("miles", "kilometers"): lambda x: x / 0.621371,
            ("celsius", "fahrenheit"): lambda x: x * 9/5 + 32,
            ("fahrenheit", "celsius"): lambda x: (x - 32) * 5/9,
            ("meters", "feet"): lambda x: x * 3.28084,
            ("feet", "meters"): lambda x: x / 3.28084,
            ("meters", "kilometers"): lambda x: x / 1000,
            ("kilometers", "meters"): lambda x: x * 1000,
            ("grams", "kg"): lambda x: x / 1000,
            ("kg", "grams"): lambda x: x * 1000,
            ("liters", "gallons"): lambda x: x * 0.264172,
            ("gallons", "liters"): lambda x: x / 0.264172,
            ("inches", "cm"): lambda x: x * 2.54,
            ("cm", "inches"): lambda x: x / 2.54,
            ("miles", "meters"): lambda x: x * 1609.34,
            ("meters", "miles"): lambda x: x / 1609.34,
            ("pounds", "kg"): lambda x: x / 2.20462,
            ("kg", "pounds"): lambda x: x * 2.20462,
            ("mb", "gb"): lambda x: x / 1024,
            ("gb", "mb"): lambda x: x * 1024,
            ("gb", "tb"): lambda x: x / 1024,
            ("tb", "gb"): lambda x: x * 1024,
        }

        key = (unit_from, unit_to)
        if key not in conversions:
            return f"Conversion from {unit_from} to {unit_to} not supported."

        result = conversions[key](value)
        return f"{value} {unit_from} = {round(result, 4)} {unit_to}"

    except Exception as e:
        return f"Unit converter error: {e}"

def currency_converter(query: str) -> str:
    import re
    try:
        CURRENCY_NAMES = {
            "dollar": "USD", "dollars": "USD", "usd": "USD",
            "rupee": "INR", "rupees": "INR", "inr": "INR",
            "euro": "EUR", "euros": "EUR", "eur": "EUR",
            "pound": "GBP", "pounds": "GBP", "gbp": "GBP",
            "yen": "JPY", "jpy": "JPY",
            "dirham": "AED", "aed": "AED",
            "singapore dollar": "SGD", "sgd": "SGD",
            "australian dollar": "AUD", "aud": "AUD",
            "canadian dollar": "CAD", "cad": "CAD",
            "swiss franc": "CHF", "chf": "CHF",
            "yuan": "CNY", "cny": "CNY", "renminbi": "CNY",
            "riyal": "SAR", "sar": "SAR",
            "won": "KRW", "krw": "KRW",
            "baht": "THB", "thb": "THB",
            "ringgit": "MYR", "myr": "MYR",
            "peso": "MXN", "mxn": "MXN",
            "real": "BRL", "brl": "BRL",
            "ruble": "RUB", "rub": "RUB",
            "lira": "TRY", "try": "TRY",
            "krona": "SEK", "sek": "SEK",
            "nok": "NOK", "krone": "NOK",
            "dkk": "DKK",
            "zloty": "PLN", "pln": "PLN",
            "taka": "BDT", "bdt": "BDT",
            "rupiah": "IDR", "idr": "IDR",
            "pkr": "PKR",
            "qar": "QAR", "qatari riyal": "QAR",
            "kwd": "KWD", "kuwaiti dinar": "KWD",
        }

        query_lower = query.lower()

        # Replace currency names with codes
        for name, code in sorted(CURRENCY_NAMES.items(), key=lambda x: -len(x[0])):
            query_lower = query_lower.replace(name, code)

        # Extract amount, from currency, to currency
        pattern = re.search(
            r'([\d,]+\.?\d*)\s*([A-Z]{3})\s+(?:to|in|into)\s+([A-Z]{3})',
            query_lower.upper()
        )

        if not pattern:
            return "Could not parse currency conversion. Use format: '100 USD to INR'"

        amount = float(pattern.group(1).replace(",", ""))
        from_currency = pattern.group(2)
        to_currency = pattern.group(3)

        # Fetch live exchange rate from frankfurter.app (free, no API key)
        url = f"https://api.frankfurter.app/latest?amount={amount}&from={from_currency}&to={to_currency}"
        response = requests.get(url, timeout=10)
        data = response.json()

        if "rates" not in data:
            return f"Could not fetch exchange rate for {from_currency} to {to_currency}"

        result = data["rates"][to_currency]
        return f"{amount:,.2f} {from_currency} = {result:,.2f} {to_currency}"

    except Exception as e:
        return f"Currency converter error: {e}"

def text_translator(query: str) -> str:
    import re
    try:
        LANGUAGE_CODES = {
            "english": "en", "hindi": "hi", "telugu": "te",
            "tamil": "ta", "kannada": "kn", "malayalam": "ml",
            "marathi": "mr", "bengali": "bn", "gujarati": "gu",
            "punjabi": "pa", "urdu": "ur", "odia": "or",
            "french": "fr", "spanish": "es", "german": "de",
            "italian": "it", "portuguese": "pt", "dutch": "nl",
            "russian": "ru", "arabic": "ar", "chinese": "zh",
            "japanese": "ja", "korean": "ko", "turkish": "tr",
            "polish": "pl", "swedish": "sv", "norwegian": "no",
            "danish": "da", "finnish": "fi", "greek": "el",
            "hebrew": "he", "thai": "th", "vietnamese": "vi",
            "indonesian": "id", "malay": "ms", "czech": "cs",
            "romanian": "ro", "hungarian": "hu", "ukrainian": "uk",
        }

        query_lower = query.lower()

        # Extract target language
        to_lang_match = re.search(
            r'(?:to|into|in)\s+([a-zA-Z]+)(?:\s|$)',
            query_lower
        )
        to_lang_name = to_lang_match.group(1).strip() if to_lang_match else "english"
        to_lang_code = LANGUAGE_CODES.get(to_lang_name, "en")

        # Extract source language if mentioned
        from_lang_match = re.search(
            r'(?:from)\s+([a-zA-Z]+)\s+(?:to|into)',
            query_lower
        )
        from_lang_code = "auto"
        if from_lang_match:
            from_lang_name = from_lang_match.group(1).strip()
            from_lang_code = LANGUAGE_CODES.get(from_lang_name, "auto")

        # Extract text to translate
        patterns = [
            r'translate\s+"([^"]+)"',
            r"translate\s+'([^']+)'",
            r'translate\s+(.+?)\s+(?:to|into|in)\s+[a-zA-Z]+',
            r'"([^"]+)"\s+(?:to|into)\s+[a-zA-Z]+',
        ]

        text_to_translate = None
        for pattern in patterns:
            match = re.search(pattern, query_lower)
            if match:
                text_to_translate = match.group(1).strip()
                break

        if not text_to_translate:
            # Use entire query minus the translation instruction
            text_to_translate = re.sub(
                r'(translate|to\s+\w+|into\s+\w+|from\s+\w+)', '', query_lower
            ).strip()

        if not text_to_translate:
            return "Please provide text to translate. Example: translate hello to french"

        # Use MyMemory API — free, no API key needed
        url = "https://api.mymemory.translated.net/get"
        lang_pair = f"{from_lang_code}|{to_lang_code}" if from_lang_code != "auto" else f"en|{to_lang_code}"
        
        params = {
            "q": text_to_translate,
            "langpair": lang_pair
        }
        
        r = requests.get(url, params=params, timeout=10)
        data = r.json()
        
        if data.get("responseStatus") == 200:
            translated = data["responseData"]["translatedText"]
            return (
                f"Original  : {text_to_translate}\n"
                f"Translated: {translated}\n"
                f"Language  : {to_lang_name.title()}"
            )
        return f"Translation failed: {data.get('responseDetails', 'Unknown error')}"

    except Exception as e:
        return f"Translator error: {e}"

def dictionary_tool(word: str) -> str:
    import re
    # Extract just the word from the query
    clean_word = re.sub(r'(define|the|word|meaning of|what is)\s*', '', word.lower()).strip()
    clean_word = clean_word.split()[0] if clean_word else word.strip()
    
    try:
        r = requests.get(
            f"https://api.dictionaryapi.dev/api/v2/entries/en/{clean_word}",
            timeout=10
        )
        data = r.json()
        if isinstance(data, list):
            meaning = data[0]["meanings"][0]
            definition = meaning["definitions"][0]["definition"]
            part_of_speech = meaning["partOfSpeech"]
            example = meaning["definitions"][0].get("example", "")
            result = f"{clean_word} ({part_of_speech}): {definition}"
            if example:
                result += f"\nExample: {example}"
            return result
        return f"No definition found for: {clean_word}"
    except Exception as e:
        return f"Dictionary error: {e}"

def weather_tool(city: str) -> str:
    import re
    try:
        clean_city = re.sub(r'(weather in|weather for|what is the weather in|how is the weather in)', '', city.lower()).strip()
        url = f"https://wttr.in/{clean_city.replace(' ', '+')}?format=j1"
        r = requests.get(url, timeout=10)
        data = r.json()
        current = data["current_condition"][0]
        area = data["nearest_area"][0]
        city_name = area["areaName"][0]["value"]
        country = area["country"][0]["value"]
        temp_c = current["temp_C"]
        feels_like = current["FeelsLikeC"]
        humidity = current["humidity"]
        wind_speed = current["windspeedKmph"]
        wind_dir = current["winddir16Point"]
        visibility = current["visibility"]
        uv_index = current["uvIndex"]
        description = current["weatherDesc"][0]["value"]
        return (
            f"Weather in {city_name}, {country}:\n"
            f"Condition    : {description}\n"
            f"Temperature : {temp_c}°C (Feels like {feels_like}°C)\n"
            f"Humidity     : {humidity}%\n"
            f"Wind        : {wind_speed} km/h {wind_dir}\n"
            f"Visibility  : {visibility} km\n"
            f"UV Index    : {uv_index}"
        )
    except Exception as e:
        try:
            r = requests.get(f"https://wttr.in/{city.replace(' ', '+')}?format=3", timeout=5)
            return r.text.strip()
        except:
            return f"Weather error: {e}"

def qr_generator(query: str) -> str:
    import re
    import qrcode
    from pathlib import Path

    try:
        clean = re.sub(
            r'(generate qr|create qr|make qr|qr code for|qr for)',
            '', query.lower()
        ).strip()

        if not clean:
            return "Please provide text or URL to generate QR code."

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4
        )
        qr.add_data(clean)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")

        upload_dir = Path(__file__).resolve().parent.parent / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)

        filename = f"qr_{int(time.time())}.png"
        filepath = upload_dir / filename
        img.save(str(filepath))

        return (
            f"QR code generated successfully!\n"
            f"Content : {clean}\n"
            f"File    : {filename}"
        )

    except Exception as e:
        return f"QR generator error: {e}"

def password_generator(query: str) -> str:
    import re
    import secrets
    import string
    try:
        query_lower = query.lower()

        # Extract length from query
        length_match = re.search(r'(\d+)\s*(?:character|char|digit|letter|length)?', query_lower)
        length = int(length_match.group(1)) if length_match else 16
        length = max(8, min(length, 64))

        # Detect what to include
        include_upper = 'no upper' not in query_lower
        include_numbers = 'no number' not in query_lower
        include_symbols = any(w in query_lower for w in [
            'symbol', 'special', 'strong', 'secure'
        ]) or 'no symbol' not in query_lower

        # Build character set
        chars = string.ascii_lowercase
        if include_upper:
            chars += string.ascii_uppercase
        if include_numbers:
            chars += string.digits
        if include_symbols:
            chars += string.punctuation

        # Generate password
        password = ''.join(secrets.choice(chars) for _ in range(length))

        # Calculate strength
        strength = "Weak"
        score = sum([
            any(c.islower() for c in password),
            any(c.isupper() for c in password),
            any(c.isdigit() for c in password),
            any(c in string.punctuation for c in password),
            length >= 12,
            length >= 16
        ])
        if score >= 5:
            strength = "Very Strong"
        elif score >= 4:
            strength = "Strong"
        elif score >= 3:
            strength = "Medium"

        return (
            f"Generated Password: `{password}`\n"
            f"Length   : {length} characters\n"
            f"Strength : {strength}\n"
            f"Contains : "
            f"{'Uppercase ' if include_upper else ''}"
            f"{'Numbers ' if include_numbers else ''}"
            f"{'Symbols' if include_symbols else ''}"
        )

    except Exception as e:
        return f"Password generator error: {e}"

def csv_analyzer(filepath: str) -> str:
    try:
        import re
        match = re.search(r'([A-Za-z]:[^\|]+\.csv)', filepath)
        actual_path = match.group(1).strip() if match else filepath.strip()
        df = pd.read_csv(actual_path)
        return f"Shape: {df.shape}\nColumns: {list(df.columns)}\nStats:\n{df.describe().to_string()[:MAX_TOOL_OUTPUT]}"
    except Exception as e:
        return f"CSV analyzer error: {e}"

def system_info(query: str) -> str:
    try:
        cpu = psutil.cpu_percent(interval=1)
        ram = psutil.virtual_memory()
        disk = psutil.disk_usage("/")
        return (
            f"OS: {platform.system()} {platform.release()}\n"
            f"CPU: {cpu}%\n"
            f"RAM: {ram.used/1e9:.1f}GB / {ram.total/1e9:.1f}GB ({ram.percent}%)\n"
            f"Disk: {disk.used/1e9:.1f}GB / {disk.total/1e9:.1f}GB ({disk.percent}%)"
        )
    except Exception as e:
        return f"System info error: {e}"

TOOL_REGISTRY = {
    "calculator":      {"fn": calculator,      "description": "Evaluates math expressions."},
    "wikipedia":       {"fn": wikipedia_tool,  "description": "Fetches Wikipedia summary."},
    "web_search":      {"fn": web_search_tool, "description": "Searches the web via DuckDuckGo."},
    "file_reader":     {"fn": file_reader,      "description": "Reads txt, csv, pdf, json files."},
    "datetime_tool":   {"fn": datetime_tool,   "description": "Returns current date/time/day."},
    "unit_converter": {"fn": unit_converter,  "description": "Converts units (km, kg, celsius, etc)."},
    "dictionary":      {"fn": dictionary_tool, "description": "Returns word definition."},
    "weather":         {"fn": weather_tool,    "description": "Returns current weather by city."},
    "csv_analyzer":   {"fn": csv_analyzer,    "description": "Analyzes CSV file statistics."},
    "system_info":     {"fn": system_info,     "description": "Returns system CPU/RAM/disk usage."},
    "currency_converter": {"fn": currency_converter, "description": "Converts currency using live exchange rates. Example: 100 USD to INR"},
    "text_translator": {"fn": text_translator, "description": "Translates text to any language. Example: translate hello to french"},
    "qr_generator": {"fn": qr_generator, "description": "Generates QR code for any text or URL. Example: generate qr code for https://google.com"},
    "password_generator": {"fn": password_generator, "description": "Generates secure random password. Example: generate a 16 character strong password"},
}

# --- app init ---
app = FastAPI(title="Autonomous AI Agent")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

static_path = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=str(static_path)), name="static")

from fastapi.staticfiles import StaticFiles as SF
upload_dir = Path(BASE_PATH) / "uploads"
upload_dir.mkdir(parents=True, exist_ok=True)
app.mount("/uploads", SF(directory=str(upload_dir)), name="uploads")

client = init_client()
config = AgentConfig()
short_term = ShortTermBuffer()
long_term = LongTermMemory()
search_history = SearchHistoryManager()
calendar = CalendarManager()

SESSIONS_PATH = Path(f"{BASE_PATH}/memory/sessions")
SESSIONS_PATH.mkdir(parents=True, exist_ok=True)

# Define and build dynamic uploads directory path anchor
UPLOAD_DIR = Path(f"{BASE_PATH}/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

current_session: dict = {"id": None, "title": None, "messages": []}

def generate_title(message: str) -> str:
    return message[:45] + "..." if len(message) > 45 else message

def session_file(session_id: str) -> Path:
    return SESSIONS_PATH / f"session_{session_id}.json"

def save_session() -> None:
    if not current_session["id"]:
        return
    session_file(current_session["id"]).write_text(
        json.dumps(current_session, indent=2), encoding="utf-8"
    )

def load_session_file(session_id: str) -> dict:
    path = session_file(session_id)
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}

def init_new_session() -> str:
    sid = datetime.now().strftime("%Y%m%d_%H%M%S")
    current_session["id"] = sid
    current_session["title"] = "New Conversation"
    current_session["messages"] = []
    current_session["created_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    save_session()
    log.info(f"New session created: {sid}")
    return sid

init_new_session()
log.info(f"Agent initialized — session: {config.session_id}")

# --- request models ---
class ChatRequest(BaseModel):
    message: str

class CalendarAddRequest(BaseModel):
    title: str
    date: str
    time: str
    description: Optional[str] = ""

# --- routes ---
@app.get("/")
def serve_index():
    return FileResponse(str(static_path / "index.html"))

@app.post("/upload")
async def upload_file(file: UploadFile = FastAPIFile(...)):
    try:
        dest = UPLOAD_DIR / file.filename
        with open(dest, "wb") as f:
            shutil.copyfileobj(file.file, f)
        log.info(f"File uploaded successfully: {dest}")
        return {"absolute_path": str(dest.resolve()).replace("\\", "/"), "filename": file.filename}
    except Exception as e:
        log.error(f"Upload processing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat")
def chat(req: ChatRequest):
    try:
        import time
        start = time.time()
        reply, intent, tool = orchestrate(
            user_input=req.message,
            client=client,
            cfg=config,
            short_term=short_term,
            long_term=long_term,
            tool_registry=TOOL_REGISTRY,
            search_history=search_history
        )
        elapsed = round(time.time() - start, 2)
        final_reply = build_response(reply, intent, tool, elapsed, config)

        if len(current_session["messages"]) == 0:
            current_session["title"] = generate_title(req.message)

        current_session["messages"].append({
            "role": "user",
            "text": req.message,
            "intent": None,
            "tool": None,
            "elapsed": None,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        current_session["messages"].append({
            "role": "agent",
            "text": final_reply,
            "intent": intent,
            "tool": tool,
            "elapsed": elapsed,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        save_session()

        return {
            "reply": final_reply,
            "intent": intent,
            "tool": tool,
            "elapsed": elapsed,
            "tokens": config.session_token_count
        }
    except Exception as e:
        log.error(f"Chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- NEW STREAMING ENDPOINT INTEGRATION ---
from fastapi.responses import StreamingResponse
import asyncio

@app.post("/chat/stream")
async def chat_stream(req: ChatRequest):
    try:
        import time
        start = time.time()
                
        # Classify intent first
        from agent.orchestrator import classify_intent
        intent, tool = classify_intent(req.message, client, config)
                
        # Only stream for chat intent
        if intent != "chat":
            reply, intent, tool = orchestrate(
                user_input=req.message,
                client=client,
                cfg=config,
                short_term=short_term,
                long_term=long_term,
                tool_registry=TOOL_REGISTRY,
                search_history=search_history
            )
            elapsed = round(time.time() - start, 2)
            final_reply = build_response(reply, intent, tool, elapsed, config)
                        
            if len(current_session["messages"]) == 0:
                current_session["title"] = generate_title(req.message)
            current_session["messages"].append({
                "role": "user", "text": req.message,
                "intent": None, "tool": None, "elapsed": None,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            current_session["messages"].append({
                "role": "agent", "text": final_reply,
                "intent": intent, "tool": tool, "elapsed": elapsed,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            save_session()
                        
            async def non_stream():
                yield f"data: {json.dumps({'type': 'metadata', 'intent': intent, 'tool': tool, 'elapsed': elapsed})}\n\n"
                yield f"data: {json.dumps({'type': 'full', 'content': final_reply})}\n\n"
                yield "data: [DONE]\n\n"
            return StreamingResponse(non_stream(), media_type="text/event-stream")
        
        # Stream for chat intent
        messages = [{"role": "system", "content": config.system_prompt}]
        messages += short_term.get_recent(6)
        messages.append({"role": "user", "content": req.message})
        
        async def generate():
            full_reply = ""
            stream = client.chat.completions.create(
                model=config.model,
                messages=messages,
                max_tokens=config.max_tokens,
                temperature=config.temperature,
                stream=True
            )
            yield f"data: {json.dumps({'type': 'metadata', 'intent': 'chat', 'tool': None})}\n\n"
                        
            for chunk in stream:
                delta = chunk.choices[0].delta
                if hasattr(delta, 'content') and delta.content:
                    full_reply += delta.content
                    yield f"data: {json.dumps({'type': 'chunk', 'content': delta.content})}\n\n"
                    await asyncio.sleep(0)
            elapsed = round(time.time() - start, 2)
            config.session_token_count += len(full_reply.split()) * 1.3
            config.total_turns += 1
            config.intent_counts["chat"] += 1
            config.response_times["chat"].append(elapsed)
            short_term.add("user", req.message)
            short_term.add("assistant", full_reply)
            long_term.store(
                memory_id=f"turn_{config.total_turns}_{int(time.time())}",
                text=f"User: {req.message}\nAssistant: {full_reply}",
                metadata={"turn": config.total_turns, "timestamp": int(time.time())}
            )
            if len(current_session["messages"]) == 0:
                current_session["title"] = generate_title(req.message)
            current_session["messages"].append({
                "role": "user", "text": req.message,
                "intent": None, "tool": None, "elapsed": None,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            current_session["messages"].append({
                "role": "agent", "text": full_reply,
                "intent": "chat", "tool": None, "elapsed": elapsed,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            save_session()
            yield f"data: {json.dumps({'type': 'done', 'elapsed': elapsed})}\n\n"
            yield "data: [DONE]\n\n"
        return StreamingResponse(generate(), media_type="text/event-stream")
    except Exception as e:
        log.error(f"Stream endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/memory")
def get_memory():
    return {
        "short_term": short_term.summary(),
        "long_term_count": long_term.count(),
        "memory_hits": config.memory_hits,
        "turns": config.total_turns
    }

@app.get("/tools")
def get_tools():
    return {
        name: meta["description"]
        for name, meta in TOOL_REGISTRY.items()
    }

@app.get("/stats")
def get_stats():
    avg_times = {
        intent: round(sum(times) / len(times), 2)
        for intent, times in config.response_times.items() if times
    }
    return {
        "session_id": config.session_id,
        "turns": config.total_turns,
        "tokens": config.session_token_count,
        "tool_calls": config.tool_call_count,
        "searches": config.search_count,
        "memory_hits": config.memory_hits,
        "intent_counts": config.intent_counts,
        "avg_response_times": avg_times
    }

@app.get("/search/history")
def get_search_history():
    return {"history": search_history.get_all()}

@app.get("/calendar")
def get_calendar():
    return {"upcoming": calendar.upcoming_events(n=10)}

@app.post("/calendar/add")
def add_calendar_event(req: CalendarAddRequest):
    result = calendar.add_event(req.title, req.date, req.time, req.description)
    return {"result": result}

@app.delete("/calendar/{event_id}")
def delete_calendar_event(event_id: str):
    result = calendar.delete_event(event_id)
    return {"result": result}

@app.post("/clear")
def clear_memory():
    short_term.clear()
    long_term.clear()
    log.info("Memory cleared via API")
    return {"status": "All memory cleared"}

@app.get("/sessions")
def list_sessions():
    sessions = []
    for f in sorted(SESSIONS_PATH.glob("session_*.json"), reverse=True):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            sessions.append({
                "id": data.get("id"),
                "title": data.get("title", "Untitled"),
                "created_at": data.get("created_at", ""),
                "message_count": len(data.get("messages", []))
            })
        except Exception:
            continue
    return {"sessions": sessions} 

@app.get("/sessions/{session_id}")
def get_session(session_id: str):
    data = load_session_file(session_id)
    if not data:
        raise HTTPException(status_code=404, detail="Session not found")
    return data

@app.post("/sessions/new")
def new_session():
    sid = init_new_session()
    short_term.clear()
    config.total_turns = 0
    config.session_token_count = 0
    config.tool_call_count = 0
    config.search_count = 0
    config.memory_hits = 0
    config.intent_counts = {
        "chat": 0, "tool": 0, "search": 0,
        "memory_recall": 0, "file_analysis": 0
    }
    config.response_times = {
        "chat": [], "tool": [], "search": [],
        "memory_recall": [], "file_analysis": []
    }
    log.info(f"New session started: {sid}")
    return {"session_id": sid}

@app.delete("/sessions/{session_id}")
def delete_session(session_id: str):
    path = session_file(session_id)
    if path.exists():
        path.unlink()
        log.info(f"Session deleted: {session_id}")
        return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="Session not found")

@app.get("/sessions/current/id")
def get_current_session_id():
    return {"session_id": current_session["id"]}

@app.get("/export")
def export_session():
    try:
        export_path = Path(EXPORT_PATH)
        export_path.parent.mkdir(parents=True, exist_ok=True)
        summary = build_session_summary(config)
        history = short_term.export()
        content = f"SESSION EXPORT\n{'='*50}\n{summary}\n\n{'='*50}\nCONVERSATION HISTORY\n{'='*50}\n{history}"
        export_path.write_text(content, encoding="utf-8")
        return FileResponse(
            path=str(export_path),
            filename=f"session_{config.session_id}.txt",
            media_type="text/plain"
        )
    except Exception as e:
        log.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
